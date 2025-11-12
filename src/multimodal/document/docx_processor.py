"""
DOCX処理
"""
import logging
import zipfile
import os
from typing import Any, Dict, List, Optional, Tuple


class DOCXProcessor:
    """DOCX処理クラス"""

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_docx_text(self, docx_path: str) -> str:
        """DOCXファイルから構造化テキストを抽出"""
        try:
            from docx import Document
        except ImportError:
            self.logger.warning("python-docxが利用できないため、簡易テキスト抽出を使用します")
            # フォールバック: ZIP内のdocument.xmlからテキスト抽出
            try:
                with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                    if 'word/document.xml' in docx_zip.namelist():
                        doc_xml = docx_zip.read('word/document.xml').decode('utf-8')
                        # XMLからテキスト部分を抽出
                        import re
                        text_parts = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', doc_xml)
                        return '\n'.join(text_parts)
            except Exception as e:
                self.logger.error(f"DOCX簡易テキスト抽出失敗: {e}")
            return ""

        try:
            doc = Document(docx_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # テーブルも処理
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            return '\n\n'.join(text_parts)

        except Exception as e:
            self.logger.error(f"DOCX構造化テキスト抽出失敗: {e}")
            return ""

    def resolve_figure_references(self, para_info: Dict[str, Any], all_paragraphs: List[Dict[str, Any]],
                                  image_metadata_list: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """図表参照を解決して対応する画像コンテンツと結びつける"""
        resolved_refs = []
        para_text = para_info['text'].lower()
        para_index = para_info['index']

        # 図表参照パターンを探す
        ref_patterns = [
            r'図\s*\d+', r'表\s*\d+', r'figure\s*\d+', r'table\s*\d+',
            r'図\s*[a-zA-Z]', r'表\s*[a-zA-Z]'
        ]

        import re
        for pattern in ref_patterns:
            matches = re.findall(pattern, para_text)
            for match in matches:
                m = re.search(r'\d+|[a-zA-Z]+', match)
                if not m:
                    continue
                ref_id = m.group()
                ref_type = '図' if '図' in match or 'figure' in match.lower() else '表'

                # 対応する画像を探す
                referenced_image = self._find_referenced_image(
                    ref_id, ref_type, para_index, all_paragraphs, image_metadata_list
                )

                if referenced_image:
                    confidence = self._calculate_reference_confidence(
                        ref_id, ref_type, para_index, referenced_image
                    )

                    resolved_refs.append({
                        'reference_id': ref_id,
                        'reference_type': ref_type,
                        'paragraph_index': para_index,
                        'image_id': referenced_image.get('id'),
                        'confidence': confidence,
                        'context': para_info['text'][:100]
                    })

        return resolved_refs

    def _find_referenced_image(self, ref_id: str, ref_type: str, current_para_idx: int,
                              all_paragraphs: List[Dict[str, Any]], image_metadata_list: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """参照された画像を見つける"""
        # 番号ベースの参照（図1、表Aなど）
        if ref_id and ref_id.replace('.', '').isdigit():
            target_num = int(ref_id.replace('.', ''))

            # 現在の段落から後方を検索（参照は通常画像の前に来る）
            search_range = all_paragraphs[max(0, current_para_idx - 10):current_para_idx + 20]

            for para in search_range:
                para_text = para['text']
                if f'{ref_type}{target_num}' in para_text:
                    # 対応する画像メタデータを探す
                    for img_meta in image_metadata_list:
                        if img_meta.get('reference_number') == target_num and img_meta.get('reference_type') == ref_type:
                            return img_meta

        return None

    def _calculate_reference_confidence(self, ref_id: str, ref_type: str, para_idx: int, target_image: Dict[str, Any]) -> float:
        """参照の信頼度を計算"""
        confidence = 0.5  # 基本信頼度

        # 番号ベースの参照は信頼度高い
        if ref_id and ref_id.replace('.', '').isdigit():
            confidence += 0.3

        # 画像との距離が近いほど信頼度高い
        image_idx = target_image.get('paragraph_index', para_idx)
        distance = abs(para_idx - image_idx)
        if distance <= 5:
            confidence += 0.2

        return min(confidence, 1.0)

    def extract_images_from_docx(self, docx_path: str) -> List[Tuple[bytes, Dict[str, Any]]]:
        """DOCXから画像を抽出（依存関係問題を回避）"""
        images: List[Tuple[bytes, Dict[str, Any]]] = []
        try:
            if not os.path.exists(docx_path):
                self.logger.warning(f"DOCXファイルが見つかりません: {docx_path}")
                return images

            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # word/media/ 内の画像ファイルを検索
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                for media_file in media_files:
                    try:
                        image_bytes = docx_zip.read(media_file)
                        filename = os.path.basename(media_file)

                        # 基本的なメタデータ
                        metadata = {
                            'filename': filename,
                            'source': 'docx_embedded',
                            'content_type': self._guess_image_type(filename),
                            'file_size': len(image_bytes),
                            'extraction_method': 'zip_extraction'
                        }

                        images.append((image_bytes, metadata))

                    except Exception as e:
                        self.logger.warning(f"画像抽出失敗 {media_file}: {e}")
                        continue

        except Exception as e:
            self.logger.error(f"DOCX画像抽出エラー: {e}")

        return images

    def _guess_image_type(self, filename: str) -> str:
        """ファイル名から画像タイプを推測"""
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            return f'image/{ext}'
        return 'application/octet-stream'

